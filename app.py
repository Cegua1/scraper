
import os
import requests
import zipfile
import streamlit as st
from bs4 import BeautifulSoup
from io import BytesIO
from PIL import Image
from docx import Document
from docx.shared import Inches

BASE_URL = "https://www.velo.com/dk/da/opslagsvaerket/page/"
HEADERS = {"User-Agent": "Mozilla/5.0"}

def get_article_links():
    links = []
    page = 1
    while True:
        res = requests.get(BASE_URL + str(page), headers=HEADERS)
        if res.status_code != 200:
            break
        soup = BeautifulSoup(res.text, "html.parser")
        articles = soup.select(".ArticleListItemstyles__StyledLink-sc-1otcs2m-0")
        if not articles:
            break
        for a in articles:
            href = a.get("href")
            if href and href not in links:
                links.append(href)
        page += 1
    return links

def download_image(img_url):
    try:
        img_data = requests.get(img_url, headers=HEADERS).content
        return BytesIO(img_data)
    except:
        return None

def scrape_article(url):
    res = requests.get(url, headers=HEADERS)
    soup = BeautifulSoup(res.text, "html.parser")

    title_tag = soup.find("h1")
    title = title_tag.get_text(strip=True) if title_tag else "Untitled"

    content_div = soup.find("div", class_="RichTextstyles__StyledWrapper-sc-mb5gzs-0")
    paragraphs = content_div.find_all("p") if content_div else []

    doc = Document()
    doc.add_heading(title, 0)

    for p in paragraphs:
        doc.add_paragraph(p.get_text(strip=True))

    images = content_div.find_all("img") if content_div else []
    for img in images:
        img_url = img.get("src")
        img_data = download_image(img_url)
        if img_data:
            try:
                image = Image.open(img_data)
                image.save("temp_img.jpg")
                doc.add_picture("temp_img.jpg", width=Inches(5))
                os.remove("temp_img.jpg")
            except:
                continue

    filename = title.replace(" ", "_").replace("/", "_") + ".docx"
    doc.save(filename)
    return filename

def zip_docs(doc_files):
    zip_buffer = BytesIO()
    with zipfile.ZipFile(zip_buffer, "w") as zipf:
        for file in doc_files:
            zipf.write(file)
            os.remove(file)
    zip_buffer.seek(0)
    return zip_buffer

st.title("Velo Blog Scraper")

if st.button("Start Scraping"):
    with st.spinner("Scraping articles..."):
        links = get_article_links()
        doc_files = []
        for link in links:
            filename = scrape_article(link)
            doc_files.append(filename)

        zip_file = zip_docs(doc_files)

    st.success(f"Scraped {len(doc_files)} articles.")
    st.download_button(
        label="Download All Articles as ZIP",
        data=zip_file,
        file_name="velo_articles.zip",
        mime="application/zip"
    )
